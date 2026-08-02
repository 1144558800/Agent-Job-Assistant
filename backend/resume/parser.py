# -*- coding: utf-8 -*-
"""
简历解析模块 - 从 PDF/Word 文件中提取文本内容
"""
from pathlib import Path
from typing import Optional
from loguru import logger


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """从 PDF 文件中提取文本"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.error("PyMuPDF 未安装，无法解析 PDF 文件")
        return None

    try:
        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        full_text = "\n".join(text_parts).strip()
        logger.info(f"PDF 解析完成: {file_path}, {len(full_text)} 字符")
        return full_text
    except Exception as e:
        logger.error(f"PDF 解析失败: {e}")
        return None


def _extract_docx_raw_xml_text(file_path: str) -> str:
    """当 python-docx 无法解析时，直接从 DOCX 的 XML 中提取所有文本"""
    import zipfile
    import xml.etree.ElementTree as ET

    try:
        with zipfile.ZipFile(file_path) as z:
            xml_content = z.read('word/document.xml')
    except Exception as e:
        logger.error(f"读取 DOCX XML 失败: {e}")
        return ""

    ns_url = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    try:
        root = ET.fromstring(xml_content)
        # 提取每个段落 <w:p> 内的所有 <w:t> 文本，拼成一行
        paragraphs = []
        for p_elem in root.iter(f'{{{ns_url}}}p'):
            # 收集该段落内所有 <w:t> 文本
            parts = []
            for t_elem in p_elem.iter(f'{{{ns_url}}}t'):
                if t_elem.text:
                    parts.append(t_elem.text)
            line = ''.join(parts).strip()
            if line:
                paragraphs.append(line)
        return "\n".join(paragraphs).strip()
    except Exception as e:
        logger.error(f"解析 DOCX XML 失败: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """从 Word (.docx) 文件中提取文本"""
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx 未安装，无法解析 Word 文件")
        return None

    try:
        doc = Document(file_path)
        text_parts = []
        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        # 提取表格中的文本（中文简历大量使用表格布局）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt:
                        text_parts.append(txt)
        full_text = "\n".join(text_parts).strip()
        # 如果 python-docx 未能提取到文本，尝试直接解析 XML
        if not full_text:
            logger.warning(f"python-docx 提取文本为空，尝试直接解析 XML: {file_path}")
            full_text = _extract_docx_raw_xml_text(file_path)
        logger.info(f"Word 解析完成: {file_path}, {len(full_text)} 字符")
        return full_text
    except Exception as e:
        logger.error(f"Word 解析失败: {e}")
        return None


def extract_text_from_txt(file_path: str) -> Optional[str]:
    """从 TXT 文件中提取文本"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read().strip()
        if text:
            logger.info(f"TXT 解析完成: {file_path}, {len(text)} 字符")
        else:
            logger.warning(f"TXT 文件内容为空: {file_path}")
        return text or None
    except UnicodeDecodeError:
        # UTF-8 解码失败，尝试 GBK 编码（部分老旧 Windows 文件用 GBK）
        try:
            with open(file_path, "r", encoding="gbk") as f:
                text = f.read().strip()
            logger.info(f"TXT(GBK) 解析完成: {file_path}, {len(text)} 字符")
            return text or None
        except Exception as e:
            logger.error(f"TXT 解析失败(GBK): {e}")
            return None
    except Exception as e:
        logger.error(f"TXT 解析失败: {e}")
        return None


def extract_resume_text(file_path: str) -> Optional[str]:
    """根据文件类型自动选择合适的解析方式"""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        return extract_text_from_docx(file_path)
    elif suffix == ".doc":
        # .doc 是旧版 Word 格式，python-docx 不支持，提示用户转格式
        logger.warning(".doc 格式暂不支持，请将文件另存为 .docx 格式后重试")
        return None
    elif suffix == ".txt":
        return extract_text_from_txt(file_path)
    else:
        logger.warning(f"不支持的文件格式: {suffix}")
        return None
