# -*- coding: utf-8 -*-
"""
简历编辑模块 - AI 润色简历，保留原始 Word 格式
核心思路：从 docx 提取带标记的文本 -> AI 润色 -> 按标记替换回原位置，保留所有格式
"""
import os
import re
import copy
from pathlib import Path
from typing import List, Tuple, Optional
from loguru import logger

from config import UPLOAD_DIR


# 段落和表格单元格的标记前缀
MARKER_PARA = "SEG_P"      # 段落: SEG_P_0, SEG_P_1, ...
MARKER_TABLE = "SEG_T"     # 表格: SEG_T_0_C_0, SEG_T_0_C_1, ... (T=表格索引, C=单元格索引)
MARKER_PATTERN = re.compile(r'\[(SEG_[PT][_\w]+)\]')


def _get_direct_t_elements(parent_elem, ns: str) -> list:
    """
    获取父元素下所有直接子 w:r 内的 w:t 元素（非递归）。
    避免递归搜索到嵌套子段落的文本。
    """
    result = []
    for r_elem in parent_elem.findall(f'{{{ns}}}r'):
        for t_elem in r_elem.findall(f'{{{ns}}}t'):
            result.append(t_elem)
    return result


def _extract_docx_segments(docx_path: str) -> Tuple[List[dict], str]:
    """
    从 docx 文件中提取所有文本段落（包括文本框、表格等），生成带标记的结构化文本。
    使用原始 XML 方式遍历所有 w:p 元素，确保覆盖文本框等复杂布局。

    返回: (segments列表, 带SEG标记的合并文本)
    """
    from docx import Document
    from lxml import etree

    NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    doc = Document(docx_path)
    segments = []

    # 通过 XML 直接查找所有 w:p 元素（包括文本框、嵌套结构中的）
    body = doc.element.body
    all_paras = body.findall(f'.//{{{NS}}}p')

    for p_idx, p_elem in enumerate(all_paras):
        # 只取当前段落的直接子 w:r 下的 w:t（非递归，避免提取嵌套子段落文本）
        t_elems = _get_direct_t_elements(p_elem, NS)
        text = ''.join(t.text or '' for t in t_elems).strip()
        if text:
            segments.append({
                "marker": f"{MARKER_PARA}_{p_idx}",
                "type": "paragraph",
                "index": p_idx,
                "text": text,
            })

    # 同时提取表格中的文本
    all_tbls = body.findall(f'.//{{{NS}}}tbl')
    cell_idx = 0
    for t_idx, tbl in enumerate(all_tbls):
        cells = tbl.findall(f'.//{{{NS}}}tc')
        for c_idx, cell_elem in enumerate(cells):
            t_elems = _get_direct_t_elements(cell_elem, NS)
            text = ''.join(t.text or '' for t in t_elems).strip()
            if text:
                segments.append({
                    "marker": f"{MARKER_TABLE}_{t_idx}_C_{cell_idx}",
                    "type": "table_cell",
                    "table_idx": t_idx,
                    "cell_idx": cell_idx,
                    "text": text,
                    "xml_element": cell_elem,
                })
                cell_idx += 1

    # 构建带标记的合并文本
    labeled_parts = []
    for seg in segments:
        labeled_parts.append(f"[{seg['marker']}] {seg['text']}")

    labeled_text = "

".join(labeled_parts)

    logger.info("docx 段提取完成(XML模式): {} 个段落段, {} 个表格段",
                sum(1 for s in segments if s["type"] == "paragraph"),
                sum(1 for s in segments if s["type"] == "table_cell"))

    return segments, labeled_text


def _ai_polish_segments(labeled_text: str, client, model: str, polish_focus: str) -> List[Tuple[str, str]]:
    """
    将带标记的文本发送给 AI 润色，返回 (marker, polished_text) 列表。

    AI 被要求严格按标记返回，不添加任何额外内容。
    """
    focus_map = {
        "project": "项目经历/项目经验部分的描述",
        "work": "工作经历/工作经验部分的描述",
        "all": "所有部分的描述（尤其项目经历和工作经历）",
    }
    focus_desc = focus_map.get(polish_focus, focus_map["project"])

    system_prompt = f"""你是一个简历文本润色器。下面是一份简历，每个文本段落用 [SEG_xxx] 标记分隔。

## 你的任务
对每个段落进行专业润色。重点优化{focus_desc}，使用 STAR 法则，突出技术深度和个人贡献。
基本信息（姓名、电话、邮箱）、教育背景（学校、专业、学历）、证书（证书名称、获得时间）等客观事实类段落保持原样。

## 输出要求（严格遵守）
1. 每个段落的润色结果独占一行，格式为: [SEG_xxx] 润色后的文本
2. 保持段落顺序不变
3. 只输出润色后的文本和标记，不要添加任何解释、说明、总结、问候语
4. 不要添加 "润色后:"、"优化版:" 之类的前缀
5. 不要添加任何不在原文中的新段落
6. 如果某段不需要润色（如姓名、电话、学校名），原样返回

## 示例
输入:
[SEG_P_0] 张三
[SEG_P_1] 负责开发后台管理系统

输出:
[SEG_P_0] 张三
[SEG_P_1] 主导企业级后台管理系统的架构设计与核心功能开发"""

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"请润色以下简历文本：

{labeled_text}"}
            ],
            temperature=0.3,
            max_tokens=8000,
        )

        raw_output = response.choices[0].message.content.strip()
        logger.info("AI 润色原始输出长度: {} 字符", len(raw_output))

        # 解析 AI 返回的结果
        polished_segments = []
        seen_markers = set()

        for line in raw_output.split('
'):
            line = line.strip()
            if not line:
                continue

            match = MARKER_PATTERN.match(line)
            if match:
                marker = match.group(1)
                # 提取标记后的文本
                rest = line[match.end():].strip()

                # 过滤明显的垃圾内容
                if rest and not _is_ai_garbage(rest):
                    if marker not in seen_markers:
                        polished_segments.append((marker, rest))
                        seen_markers.add(marker)

        logger.info("AI 解析后得到 {} 个润色段", len(polished_segments))
        return polished_segments

    except Exception as e:
        logger.error("AI 润色失败: {}", e)
        raise


def _is_ai_garbage(text: str) -> bool:
    """判断 AI 返回的文本是否为无用的解释性内容"""
    garbage_patterns = [
        r'^好的[，,]我理解',
        r'^好的[，,]已',
        r'^根据您',
        r'^以下是',
        r'^---+\s*$',
        r'^\*\*说明\*\*',
        r'^\*\*注意\*\*',
        r'^润色后',
        r'^优化版',
        r'^【润色',
        r'^---',
        r'^===',
        r'^如您有',
        r'^欢迎',
        r'^如需',
        r'^因原始简历中',
        r'^已将重复',
        r'^\*\*输出格式',
        r'^##\s',
        r'^#{1,3}\s',
    ]
    for pat in garbage_patterns:
        if re.search(pat, text):
            return True
    return False


def _apply_polish_to_docx(docx_path: str, output_path: str, polished_segments: List[Tuple[str, str]]):
    """
    将润色后的文本写回 docx，保留原始格式。
    通过 XML 级别操作，直接修改 w:t 元素中的文本。
    """
    from docx import Document
    from lxml import etree

    NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # 先复制原始文件
    import shutil
    shutil.copy2(docx_path, output_path)

    doc = Document(output_path)

    # 构建 marker -> polished_text 的映射
    polish_map = {marker: text for marker, text in polished_segments}

    body = doc.element.body
    replaced_count = 0

    # 替换段落中的文本（遍历所有 w:p 元素）
    all_paras = body.findall(f'.//{{{NS}}}p')
    for p_idx, p_elem in enumerate(all_paras):
        marker = f"{MARKER_PARA}_{p_idx}"
        if marker in polish_map:
            new_text = polish_map[marker]
            t_elems = _get_direct_t_elements(p_elem, NS)
            if t_elems:
                old_text = ''.join(t.text or '' for t in t_elems)
                if old_text.strip() and new_text != old_text:
                    # 将新文本写入第一个 w:t，清空其余
                    t_elems[0].text = new_text
                    for t in t_elems[1:]:
                        t.text = ''
                    replaced_count += 1

    # 替换表格单元格中的文本
    all_tbls = body.findall(f'.//{{{NS}}}tbl')
    cell_idx = 0
    for t_idx, tbl in enumerate(all_tbls):
        cells = tbl.findall(f'.//{{{NS}}}tc')
        for c_idx, cell_elem in enumerate(cells):
            marker = f"{MARKER_TABLE}_{t_idx}_C_{cell_idx}"
            if marker in polish_map:
                new_text = polish_map[marker]
                t_elems = _get_direct_t_elements(cell_elem, NS)
                if t_elems:
                    old_text = ''.join(t.text or '' for t in t_elems)
                    if old_text.strip() and new_text != old_text:
                        t_elems[0].text = new_text
                        for t in t_elems[1:]:
                            t.text = ''
                        replaced_count += 1
            cell_idx += 1

    doc.save(output_path)
    logger.info("docx XML级文本替换完成: {} 个段被润色, 保存至 {}", replaced_count, output_path)


def _replace_paragraph_text(para, new_text: str):
    """
    替换段落中的文本，保留原格式。
    取第一个 run 的格式属性，创建新 run 并写入新文本。
    """
    from docx.oxml.ns import qn

    if not para.runs:
        # 没有 run，直接添加
        run = para.add_run(new_text)
        return

    # 保存第一个 run 的格式
    first_run = para.runs[0]
    font = first_run.font
    saved_format = {
        "bold": font.bold,
        "italic": font.italic,
        "underline": font.underline,
        "size": font.size,
        "color_rgb": font.color.rgb if font.color and font.color.rgb else None,
        "name": font.name,
    }
    # 保存东亚字体
    saved_east_asian = None
    rPr = first_run._element.rPr
    if rPr is not None:
        for child in rPr:
            if child.tag == qn('w:rFonts'):
                saved_east_asian = child.get(qn('w:eastAsia'))
                break

    # 清空所有 run
    for run in para.runs:
        run.text = ""

    # 写入新文本到第一个 run
    first_run.text = new_text

    # 恢复格式
    _restore_run_format(first_run, saved_format, saved_east_asian)


def _replace_cell_text(cell, new_text: str):
    """替换表格单元格中的文本，保留原格式"""
    # 取第一个段落
    if cell.paragraphs:
        _replace_paragraph_text(cell.paragraphs[0], new_text)
    else:
        # 没有段落，创建一个
        para = cell.add_paragraph()
        para.add_run(new_text)


def _restore_run_format(run, saved_format: dict, east_asian: Optional[str]):
    """恢复 run 的格式属性"""
    from docx.oxml.ns import qn

    if saved_format.get("bold") is not None:
        run.bold = saved_format["bold"]
    if saved_format.get("italic") is not None:
        run.italic = saved_format["italic"]
    if saved_format.get("underline") is not None:
        run.underline = saved_format["underline"]
    if saved_format.get("size") is not None:
        run.font.size = saved_format["size"]
    if saved_format.get("color_rgb") is not None:
        run.font.color.rgb = saved_format["color_rgb"]
    if saved_format.get("name") is not None:
        run.font.name = saved_format["name"]

    # 恢复东亚字体
    if east_asian:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            from lxml import etree
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:eastAsia'), east_asian)


def polish_docx_inplace(input_path: str, output_path: str, client, model: str, polish_focus: str = "project") -> str:
    """
    原地润色 docx 简历文件，保留所有原始格式（字体、颜色、表格、排版等）。

    参数:
        input_path: 原始 docx 文件路径
        output_path: 输出文件路径
        client: OpenAI 兼容的 AI 客户端
        model: 模型名称
        polish_focus: "project" / "work" / "all"

    返回: output_path
    """
    logger.info("开始原地润色 docx: {} -> {}", input_path, output_path)

    # 步骤1: 从 docx 提取带标记的文本段
    segments, labeled_text = _extract_docx_segments(input_path)
    if not segments:
        logger.warning("docx 中未找到任何文本")
        return output_path

    # 步骤2: AI 润色
    polished_segments = _ai_polish_segments(labeled_text, client, model, polish_focus)
    if not polished_segments:
        logger.warning("AI 未返回有效润色结果，保留原始文件")
        import shutil
        shutil.copy2(input_path, output_path)
        return output_path

    # 步骤3: 将润色结果写回 docx（保留格式）
    _apply_polish_to_docx(input_path, output_path, polished_segments)

    logger.info("原地润色完成: {}", output_path)
    return output_path


# ========== 兼容旧接口（纯文本润色 + 生成新 docx）==========

def polish_resume_with_ai(resume_text: str, client, model: str, polish_focus: str = "project") -> str:
    """
    使用 AI 润色纯文本简历内容（不保留原始格式，仅用于纯文本输入）。
    """
    logger.warning("polish_resume_with_ai 为纯文本模式，不保留格式。推荐使用 polish_docx_inplace() 处理 docx 文件。")
    return _ai_polish_plain_text(resume_text, client, model, polish_focus)


def _ai_polish_plain_text(resume_text: str, client, model: str, polish_focus: str) -> str:
    """纯文本 AI 润色 """
    focus_map = {
        "project": "重点优化【项目经历/项目经验】部分，使用STAR法则重写每个项目描述，加入量化数据，突出技术难点和个人贡献。同时适度润色其他部分。",
        "work": "重点优化【工作经历】部分，将描述改写得更专业，突出职责、成果和影响力。",
        "all": "全面润色整份简历，优化所有模块的表达方式，使其更专业、更有说服力。",
    }

    system_prompt = f"""你是简历优化专家。{focus_map.get(polish_focus, focus_map['project'])}

核心原则:
1. 不编造任何经历、技术、数据
2. 不修改客观信息（姓名、联系方式、学校、专业、学历、证书名称等）
3. STAR法则优化项目描述
4. 输出完整的润色后简历全文

输出要求: 只输出润色后的简历全文，不要添加解释、说明、问候语。"""

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"请润色以下简历：

{resume_text}"}
            ],
            temperature=0.3,
            max_tokens=8000,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logger.error("AI 纯文润色失败: {}", e)
        raise


def generate_polished_docx(polished_text: str, output_path: str) -> str:
    """生成格式化的 Word 文件（仅用于纯文本输入场景）"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    known_sections = [
        "基本信息", "个人信息", "求职意向", "教育背景", "教育经历",
        "工作经历", "工作经验", "项目经历", "项目经验",
        "专业技能", "技术技能", "个人技能", "技能特长",
        "自我评价", "个人优势", "证书资质", "语言能力",
        "实习经历", "校园经历", "获奖情况", "荣誉奖项",
    ]

    for line in polished_text.split('
'):
        line = line.strip()
        if not line:
            continue

        is_list = re.match(r'^(\s*[-*•]\s+|\s*\d+[.、]\s+)(.*)', line)

        if line in known_sections:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        elif is_list:
            content = re.sub(r'^\s*[-*•\d+.,、]\s+', '', line).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            run = p.add_run(content)
            run.font.size = Pt(11)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(line)
            run.font.size = Pt(11)

    doc.save(output_path)
    logger.info("纯文本 Word 文件已生成: {}", output_path)
    return output_path


def generate_polished_pdf(polished_text: str, output_path: str) -> str:
    """生成 PDF 文件（仅用于纯文本输入场景）"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    _register_chinese_font()

    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            topMargin=1.5 * cm, bottomMargin=1.5 * cm,
                            leftMargin=2.0 * cm, rightMargin=2.0 * cm)

    body_style = ParagraphStyle('Body', fontName='ChineseFont', fontSize=11, leading=20, spaceAfter=6)
    heading_style = ParagraphStyle('Heading', fontName='ChineseFont', fontSize=15, leading=24,
                                   spaceBefore=16, spaceAfter=8, textColor=HexColor('#1A56DB'))

    story = []
    for line in polished_text.split('
'):
        line = line.strip()
        if not line:
            continue
        if line in ["基本信息", "个人信息", "求职意向", "教育背景", "教育经历",
                     "工作经历", "工作经验", "项目经历", "项目经验",
                     "专业技能", "技术技能", "个人技能", "技能特长",
                     "自我评价", "个人优势", "证书资质", "语言能力"]:
            story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor('#CCCCCC')))
            story.append(Spacer(1, 4))
            story.append(Paragraph(line, heading_style))
        else:
            story.append(Paragraph(line, body_style))

    doc.build(story)
    logger.info("纯文本 PDF 文件已生成: {}", output_path)
    return output_path


def _register_chinese_font():
    """注册中文字体"""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_candidates = [
        ("C:/Windows/Fonts/msyh.ttc", "微软雅黑"),
        ("C:/Windows/Fonts/msyhbd.ttc", "微软雅黑"),
        ("C:/Windows/Fonts/simsun.ttc", "宋体"),
        ("C:/Windows/Fonts/simhei.ttf", "黑体"),
    ]
    for font_path, font_name in font_candidates:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                logger.info("PDF 中文字体注册成功: {}", font_name)
                return
            except Exception:
                continue
    logger.warning("未找到中文字体，PDF 中文可能无法显示")


def detect_resume_language(resume_text: str) -> str:
    """检测简历语言类型"""
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', resume_text))
    english_chars = len(re.findall(r'[a-zA-Z]', resume_text))
    total = chinese_chars + english_chars
    if total == 0:
        return "unknown"
    ratio = chinese_chars / total
    if ratio > 0.6:
        return "chinese"
    elif ratio < 0.2:
        return "english"
    return "mixed"
